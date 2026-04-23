# backend/app/exporter.py
import io
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.cell.rich_text import CellRichText, TextBlock
from openpyxl.cell.text import InlineFont
from docx import Document
from .models import SeatingPlan, RoomPlan, SeatAssignment

_HEADERS = ["Tisch", "Platz", "Nachname", "Vorname", "Klasse", "Fach", "Dauer (min)", "Hilfsmittel", "Lehrkraft"]
_ROOMS = [("room_a", "Raum A"), ("room_b", "Raum B"), ("room_c", "Raum C")]

# --- Grid layout constants ---
_GRID_COLS = 8   # 4 desks × 2 seats per row
_GRID_ROWS = 4   # 4 desk rows
_TITLE_ROW = 1
_PULT_ROW = 2
_FIRST_SEAT_ROW = 4  # row 3 stays blank as visual spacer
_SEAT_ROW_HEIGHT = 85

_THIN = Side(style="thin", color="999999")
_THICK = Side(style="medium", color="333333")


def _style_title_row(ws, text: str) -> None:
    ws.cell(_TITLE_ROW, 1).value = text
    ws.merge_cells(start_row=_TITLE_ROW, start_column=1, end_row=_TITLE_ROW, end_column=_GRID_COLS)
    c = ws.cell(_TITLE_ROW, 1)
    c.font = Font(bold=True, color="FFFFFF", size=14)
    c.fill = PatternFill("solid", fgColor="D97706")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[_TITLE_ROW].height = 26


def _style_pult_row(ws) -> None:
    ws.cell(_PULT_ROW, 1).value = "Lehrpult"
    ws.merge_cells(start_row=_PULT_ROW, start_column=1, end_row=_PULT_ROW, end_column=_GRID_COLS)
    c = ws.cell(_PULT_ROW, 1)
    c.font = Font(italic=True, color="666666", size=10)
    c.fill = PatternFill("solid", fgColor="EEEEEE")
    c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[_PULT_ROW].height = 18


def _seat_cell_value(assignment: SeatAssignment | None, desk: int, seat: int) -> CellRichText | str:
    label = f"T{desk}.S{seat}"
    if assignment is None:
        return label
    st = assignment.student
    en = assignment.entry
    font_label = InlineFont(sz=8, color="888888")
    font_name = InlineFont(sz=11, b=True)
    font_body = InlineFont(sz=9)
    font_aids = InlineFont(sz=9, i=True)
    has_aids = bool(en.aids.strip())
    teacher_text = f"{en.teacher}\n" if has_aids else en.teacher
    runs: list = [
        TextBlock(font_label, f"{label}\n"),
        TextBlock(font_name, f"{st.last_name}, {st.first_name}\n"),
        TextBlock(font_body, f"{st.class_name}\n"),
        TextBlock(font_body, f"{en.subject} · {en.duration_minutes} min\n"),
        TextBlock(font_body, teacher_text),
    ]
    if has_aids:
        runs.append(TextBlock(font_aids, en.aids.strip()))
    return CellRichText(*runs)


def _style_seat_cell(cell, filled: bool, desk_col_index: int) -> None:
    # desk_col_index 0..3 → which desk column within the row
    # Seats 1 and 2 of the same desk sit in columns (2*desk_col_index+1, +2).
    # Between desks we want a THICK right border on seat 2 (when desk_col_index < 3).
    # Between desk rows we want THICK top/bottom borders.
    # Inside a desk, between seat 1 and seat 2, THIN vertical border.
    col_within_desk = (cell.column - 1) % 2  # 0 = seat1, 1 = seat2
    left = _THICK if col_within_desk == 0 else _THIN
    right = _THICK if col_within_desk == 1 else _THIN
    # Outer top/bottom of each desk row is THICK
    top = _THICK
    bottom = _THICK
    cell.border = Border(left=left, right=right, top=top, bottom=bottom)
    cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
    if not filled:
        cell.font = Font(size=9, color="AAAAAA")


def _render_room_grid(ws, room_plan: RoomPlan, title: str) -> None:
    _style_title_row(ws, f"{title} — Nachschreiber")
    _style_pult_row(ws)
    # Row 3 is blank spacer
    ws.row_dimensions[3].height = 8

    assignment_map = {(a.desk, a.seat): a for a in room_plan.assignments}

    for row_idx in range(_GRID_ROWS):
        excel_row = _FIRST_SEAT_ROW + row_idx
        ws.row_dimensions[excel_row].height = _SEAT_ROW_HEIGHT
        for desk_col in range(4):  # 4 desks per row
            desk_number = row_idx * 4 + desk_col + 1  # 1..16
            for seat_in_desk in range(2):
                seat_number = seat_in_desk + 1  # 1 or 2
                excel_col = desk_col * 2 + seat_in_desk + 1
                cell = ws.cell(excel_row, excel_col)
                assignment = assignment_map.get((desk_number, seat_number))
                cell.value = _seat_cell_value(assignment, desk_number, seat_number)
                _style_seat_cell(cell, filled=assignment is not None, desk_col_index=desk_col)

    # Uniform column widths for all 8 seat columns
    for col_i in range(1, _GRID_COLS + 1):
        ws.column_dimensions[get_column_letter(col_i)].width = 20

    # Print setup: landscape, fit to 1 page
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_options.horizontalCentered = True


def build_excel(plan: SeatingPlan) -> bytes:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for attr, title in _ROOMS:
        ws = wb.create_sheet(title)
        _render_room_grid(ws, getattr(plan, attr), title)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _add_room_table(doc: Document, room_plan: RoomPlan, title: str) -> None:
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(_HEADERS):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True

    assignment_map = {(a.desk, a.seat): a for a in room_plan.assignments}
    for desk in range(1, 17):
        for seat in range(1, 3):
            a = assignment_map.get((desk, seat))
            row = table.add_row().cells
            row[0].text = str(desk)
            row[1].text = str(seat)
            if a:
                row[2].text = a.student.last_name
                row[3].text = a.student.first_name
                row[4].text = a.student.class_name
                row[5].text = a.entry.subject
                row[6].text = str(a.entry.duration_minutes)
                row[7].text = a.entry.aids
                row[8].text = a.entry.teacher


def build_word(plan: SeatingPlan) -> bytes:
    doc = Document()
    doc.core_properties.title = "Nachschreiber Sitzplan"
    for i, (attr, title) in enumerate(_ROOMS):
        if i > 0:
            doc.add_page_break()
        _add_room_table(doc, getattr(plan, attr), title)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
